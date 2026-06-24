#!/usr/bin/env python3
from pathlib import Path
import csv

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "docs" / "report_10000"
METRICS = REPORT_DIR / "derived_process_metrics.csv"
CHECKSUMS = ROOT / "results" / "report_10000_process_sweep_checksum_summary.csv"
FIGURES = REPORT_DIR / "figures"
GROUP_INFO = ROOT / "config" / "group_info.txt"
OUT = ROOT / "docs" / "BaoCao_DeTai1_NhanMaTran_MPI.docx"


def read_rows(path):
    if not path.exists():
        return []
    with path.open(newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def group_info():
    if not GROUP_INFO.exists():
        return "Nhóm thực hiện: [Điền tên nhóm]"
    names = []
    with GROUP_INFO.open(encoding="utf-8") as f:
        for raw in f:
            line = raw.strip()
            if line and not line.startswith("#"):
                names.append(line)
    return "Nhóm thực hiện: " + ("; ".join(names) if names else "[Điền tên nhóm]")


def shade(cell, fill="EAF1F8"):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
        shade(table.rows[0].cells[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()
    return table


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if p.runs:
        p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p


def para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.08
    return p


def bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    return p


def add_chart(doc, filename, caption):
    path = FIGURES / filename
    if not path.exists():
        para(doc, f"Thiếu hình: {path}")
        return
    doc.add_picture(str(path), width=Inches(6.3))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)


def fmt(value, digits=3):
    try:
        return f"{float(value):.{digits}f}"
    except (TypeError, ValueError):
        return str(value)


def build_doc():
    metrics = read_rows(METRICS)
    checksum_rows = read_rows(CHECKSUMS)

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BÁO CÁO CUỐI KỲ\nĐỀ TÀI 1: SONG SONG HÓA NHÂN MA TRẬN KÍCH THƯỚC LỚN")
    run.bold = True
    run.font.size = Pt(17)
    run.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    subtitle = doc.add_paragraph("Học phần: Tính toán hiệu năng cao (High Performance Computing)")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info = doc.add_paragraph(group_info())
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading(doc, "1. Giới thiệu và cơ sở lý thuyết")
    para(doc, "Đề tài xây dựng chương trình nhân ma trận lớn bằng C++ và Open MPI, gồm bản tuần tự và các bản song song để đánh giá hiệu năng khi thay đổi số process và số node.")
    para(doc, "Bài toán tổng quát là A(M x K) * B(K x N) = C(M x N). Mỗi phần tử C[i,j] được tính bằng tổng A[i,t] * B[t,j] với t chạy từ 0 đến K-1. Độ phức tạp tuần tự là O(MKN), tương ứng xấp xỉ 2*M*K*N FLOP.")
    para(doc, "Benchmark chính dùng shape 10000x10000x10000, tương ứng khoảng 2 x 10^12 FLOP. Dữ liệu A/B là số nguyên ngẫu nhiên trong [0,1000] nhưng lưu bằng double để giữ cấu trúc thuật toán và checksum.")

    add_table(doc, ["Variant", "Vai trò"], [
        ["seq", "Bản tuần tự cơ sở."],
        ["seq_tiled", "Bản tuần tự có tile để cải thiện locality cache."],
        ["mpi", "MPI chia hàng A, broadcast toàn bộ B, gather C theo hàng."],
        ["mpi_tiled", "MPI chia hàng kết hợp local multiply dạng tile."],
        ["mpi_2d", "MPI chia lưới 2D theo ý tưởng matrix-v2: mỗi rank nhận block hàng A và block cột B đã transpose."],
    ])

    heading(doc, "2. Thiết kế MPI/Open MPI")
    para(doc, "Biến thể báo cáo tập trung là mpi_2d. Chương trình chọn layout 2D từ các ước số của P theo mục tiêu cân bằng lưới và phù hợp tỉ lệ M/N, cùng tinh thần với MPI_Dims_create. Log tổng hợp xác nhận p9 là 3x3, p12 là 3x4 và p44 là 4x11. Rank 0 sinh dữ liệu theo seed cố định, pack A theo block hàng, pack B theo block cột ở dạng transpose, gửi dữ liệu cho từng rank và reduce checksum sau khi tính.")
    para(doc, "Luồng rút gọn: rank 0 sinh A/B, chọn grid 2D, pack A_local và B_local_T cho từng rank; mỗi rank tính C_local; checksum cục bộ được reduce về rank 0. Với benchmark lớn, chương trình không gather toàn bộ C nếu không lưu sample/full result. Cách này giảm bộ nhớ và I/O, nhưng rank 0 vẫn là điểm tập trung pack/send nên có thể thành bottleneck khi số process quá cao.")
    add_table(doc, ["Thành phần", "Thiết kế"], [
        ["Dữ liệu", "Ma trận chữ nhật tổng quát M,K,N; random deterministic theo seed."],
        ["Phân phối", "2D grid chia hàng của C theo M và cột của C theo N."],
        ["Local compute", "Mỗi rank nhân A_local với B_local^T để tính C_local."],
        ["Đúng đắn", "Checksum hash giống nhau giữa các cấu hình process."],
        ["Đo thời gian", "Runtime lấy trong phần chạy MPI, không tính thao tác shell/chuẩn bị thủ công."],
    ])

    heading(doc, "3. Môi trường và phương pháp đo")
    para(doc, "Thực nghiệm cuối chạy trên cụm 3 máy WSL trong cùng mạng LAN. Tổng logical core là 44: node1 có 12 core, node2 có 16 core và node3 có 16 core. Open MPI chạy qua SSH/TCP, hostfile ép số slot theo từng cấu hình để biết chính xác mỗi node đóng góp bao nhiêu process.")
    para(doc, "Các cấu hình đo chính gồm p4, p9, p12, p16, p25, p36 và p44. p9 và p16 được chạy 3 lần để có trung bình và độ lệch chuẩn; các cấu hình còn lại chạy 1 lần do dữ liệu rất lớn.")
    para(doc, "Speedup và efficiency trong báo cáo là scaling tương đối theo p4 vì benchmark chính không chạy được bản tuần tự trong khung thời gian thực nghiệm. Đây là hạn chế được nêu rõ, không được diễn giải như speedup chuẩn T(1)/T(p).")

    heading(doc, "4. Kết quả thực nghiệm")
    if metrics:
        table_rows = []
        for r in metrics:
            pcount = int(float(r["processes"]))
            if pcount == 4:
                grid, mapping = "2x2", "2/1/1"
            elif pcount == 9:
                grid, mapping = "3x3", "3/3/3"
            elif pcount == 12:
                grid, mapping = "3x4", "4/4/4"
            elif pcount == 16:
                grid, mapping = "4x4", "4/6/6"
            elif pcount == 25:
                grid, mapping = "5x5", "7/9/9"
            elif pcount == 36:
                grid, mapping = "6x6", "10/13/13"
            else:
                grid, mapping = "4x11", "12/16/16"
            runs = "3" if pcount in (9, 16) else "1"
            table_rows.append([
                str(pcount),
                grid,
                runs,
                fmt(r["time_min_sec"]),
                fmt(r["time_avg_sec"]),
                fmt(r["time_std_sec"]),
                fmt(r["best_gflops"]),
                mapping,
            ])
        add_table(doc, ["P", "Grid", "Runs", "Min(s)", "Avg(s)", "Std(s)", "GFLOP/s", "Mapping"], table_rows)

    para(doc, "Trong các cấu hình đã đo, p9 có best/min runtime tốt nhất với 144.972 giây và 13.796 GFLOP/s. p12 đứng thứ hai với 148.039 giây, chỉ chậm hơn khoảng 2.12%. Vì p9 chạy 3 lần còn p12 mới chạy 1 lần, nên cần chạy lặp p12 trong điều kiện mạng tương tự nếu muốn kết luận ổn định hơn. Ngược lại, p44 dùng toàn bộ 44 logical core nhưng chậm hơn p9 hơn 2 lần do overhead pack/send và số block/thông điệp tăng mạnh.")
    add_chart(doc, "runtime_by_process.png", "Hình 1. Runtime tốt nhất theo số process MPI.")
    add_chart(doc, "gflops_by_process.png", "Hình 2. Thông lượng GFLOP/s theo số process MPI.")
    add_chart(doc, "speedup_vs_p4.png", "Hình 3. Speedup tương đối theo mốc p4.")
    add_chart(doc, "efficiency_vs_p4.png", "Hình 4. Efficiency tương đối theo mốc p4.")

    heading(doc, "5. Kiểm tra đúng đắn và phân tích")
    if checksum_rows:
        add_table(doc, ["Variant", "Shape", "P", "Nodes", "Checksum status"], [
            [
                r["variant"],
                f'{r["m"]}x{r["k"]}x{r["n"]}',
                r["processes"],
                r["nodes"],
                r["status"],
            ]
            for r in checksum_rows
        ])
    para(doc, "Tất cả cấu hình trong sweep có cùng checksum hash c14405814301863a. Điều này chứng minh khác biệt runtime không đến từ sai dữ liệu đầu vào, bỏ sót rank hoặc sai layout 2D.")
    para(doc, "p9 có best/min runtime tốt nhất vì nằm gần điểm cân bằng của hệ thống: đủ process để khai thác cả 3 node nhưng chưa tạo quá nhiều block và thông điệp. p12 rất gần p9, còn từ p16 trở lên overhead giao tiếp bắt đầu lớn hơn lợi ích compute.")
    para(doc, "p44 không nhanh nhất vì layout 4x11 chia nhiều cột, làm số block B, số buffer pack và số lần gửi tăng rõ rệt. Với thiết kế hiện tại, rank 0 tự pack/gửi tuần tự nên không khai thác hết lợi thế của 44 core.")

    heading(doc, "6. Hạn chế, hướng tối ưu và kết luận")
    bullet(doc, "Hạn chế chính: rank 0 vẫn tập trung sinh dữ liệu, pack và gửi dữ liệu cho các rank.")
    bullet(doc, "Môi trường WSL qua LAN/Wi-Fi không ổn định như cluster HPC thật; các thử nghiệm 15000/20000 bị Open MPI TCP timeout nên không dùng làm kết quả hiệu năng chính.")
    bullet(doc, "Benchmark chính chưa có sweep đầy đủ 1/2/3 node và speedup chuẩn T(1)/T(p); báo cáo dùng scaling tương đối theo p4 và nêu rõ hạn chế này.")
    bullet(doc, "Hướng tối ưu tiếp theo: dùng SUMMA hoặc broadcast panel theo communicator 2D, non-blocking MPI và sinh dữ liệu phân tán trực tiếp trên từng rank.")
    para(doc, "Kết luận: chương trình đã có bản tuần tự, bản MPI chia hàng và bản mpi_2d theo lưới 2D. Trên benchmark chính, p9 có best/min runtime tốt nhất trong các cấu hình đã đo; p12 gần tương đương nhưng nên chạy lặp thêm để kết luận ổn định hơn; p44 chậm hơn do bottleneck giao tiếp. Kết quả checksum xác nhận tính đúng và tính tái lập.")

    heading(doc, "Phụ lục: nguồn dữ liệu")
    bullet(doc, "docs/report_10000/BaoCao_MPI_NhanMaTran_10000.pdf")
    bullet(doc, "results/report_10000_process_sweep_raw.csv")
    bullet(doc, "results/report_10000_process_sweep_summary.csv")
    bullet(doc, "results/report_10000_process_sweep_checksum_summary.csv")
    bullet(doc, "docs/report_10000/derived_process_metrics.csv")

    heading(doc, "Tài liệu tham khảo")
    bullet(doc, "Open MPI Project. Open MPI Documentation. https://docs.open-mpi.org/")
    bullet(doc, "Open MPI Project. MPI_Dims_create manual page. https://docs.open-mpi.org/en/v5.0.3/man-openmpi/man3/MPI_Dims_create.3.html")
    bullet(doc, "Robert A. van de Geijn và Jerrell Watts. SUMMA: Scalable Universal Matrix Multiplication Algorithm. https://www.netlib.org/lapack/lawnspdf/lawn96.pdf")
    bullet(doc, "AMATH 483/583 High-Performance Scientific Computing. Lecture 20: Cannon's Algorithm. https://amath583.github.io/sp22/units/L20.html")
    bullet(doc, "Đề bài báo cáo cuối kỳ học phần Tính toán hiệu năng cao, 2026.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
